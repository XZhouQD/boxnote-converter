'''
Modified from https://github.com/pqzx/html2docx
Original License: MIT LICENSE
Removed some configurable options for simplification.
Added support for some other tags for boxnote-converter
'''

import re
import os
from urllib.parse import urlparse
from html.parser import HTMLParser

import docx, docx.table
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from bs4 import BeautifulSoup

# values in inches
INDENT = 0.25
LIST_INDENT = 0.5
BLOCKQUOTE_INDENT = 0.5
MAX_INDENT = 5.5

# Style to use with tables. By default no style is used.
DEFAULT_TABLE_STYLE = None

# Style to use with paragraphs. By default no style is used.
DEFAULT_PARAGRAPH_STYLE = None

def remove_last_occurence(ls, x):
    if x in ls:
        ls.pop(len(ls) - ls[::-1].index(x) - 1)

def remove_whitespace(string, leading=False, trailing=False):
    # Remove any leading new line characters along with any surrounding white space
    string = re.sub(r'^\s*\n+', '', string) if leading else string

    # Remove any trailing new line characters along with any surrounding white space
    string = re.sub(r'\n+\s*$', '', string) if trailing else string

    # Replace new line characters and absorb any surrounding space.
    #string = re.sub(r'\s*\n\s*', ' ', string)
    return re.sub(r'\s+', ' ', string)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

font_styles = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
}

font_names = {
    'code': 'Courier',
    'pre': 'Courier',
}

styles = {
    'LIST_BULLET': 'List Bullet',
    'LIST_NUMBER': 'List Number',
}

class HtmlToDocx(HTMLParser):

    def __init__(self):
        super().__init__()
        self.table_row_selectors = [
            'table > tr',
            'table > thead > tr',
            'table > tbody > tr',
            'table > tfoot > tr'
        ]
        self.table_style = DEFAULT_TABLE_STYLE
        self.paragraph_style = DEFAULT_PARAGRAPH_STYLE

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
            'blockquote': []
        }
        self.doc = document if document else Document()
        self.document = self.doc
        self.paragraph = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0
        self.blockquote = False
        self.style = False

    def copy_settings_from(self, other):
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def get_cell_html(self, soup):
        return ' '.join([str(i) for i in soup.contents])

    def add_styles_to_paragraph(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin = int(float(re.sub(r'[a-z]+', '', margin)))
            if units == 'px':
                self.paragraph.paragraph_format.left_indent = Inches(min(margin // 10 * INDENT, MAX_INDENT))

    def add_styles_to_run(self, style):
        if 'color' in style:
            if 'rgb' in style['color']:
                color = re.sub(r'[a-z()]+', '', style['color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['color']:
                color = style['color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
            self.run.font.color.rgb = RGBColor(*colors)
        
        if 'font-size' in style:
            if 'em' in style['font-size']:
                size = round(float(re.sub(r'[a-z]+', '', style['font-size'])) * 12)
                self.run.font.size = Pt(size)
            elif 'px' in style['font-size']:
                size = round(float(re.sub(r'[a-z]+', '', style['font-size'])) * 0.75)
                self.run.font.size = Pt(size)
            elif 'pt' in style['font-size']:
                size = round(float(re.sub(r'[a-z]+', '', style['font-size'])))
                self.run.font.size = Pt(size)

        if 'background-color' in style:
            if 'rgb' in style['background-color']:
                color = color = re.sub(r'[a-z()]+', '', style['background-color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['background-color']:
                color = style['background-color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), f"{colors[0]:02x}{colors[1]:02x}{colors[2]:02x}")
            self.run.font.size = Pt(11) if not self.run.font.size else self.run.font.size
            self.run._r.rPr.append(shd)

    def apply_paragraph_style(self, style=None):
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(f"Unable to apply style {self.paragraph_style}.") from e

    def parse_dict_string(self, string, separator=';'):
        new_string = string.replace(" ", '').split(separator)
        string_dict = dict([x.split(':') for x in new_string if ':' in x])
        return string_dict

    def handle_li(self):
        list_depth = len(self.tags['list'])
        if list_depth:
            list_type = self.tags['list'][-1]
        else:
            list_type = 'ul'

        if list_type == 'ol':
            list_style = styles['LIST_NUMBER']
        else:
            list_style = styles['LIST_BULLET']

        self.paragraph = self.doc.add_paragraph(style=list_style)            
        self.paragraph.paragraph_format.left_indent = Inches(min(list_depth * LIST_INDENT, MAX_INDENT))
        self.paragraph.paragraph_format.line_spacing = 1

    def add_image_to_cell(self, cell, image):
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image)

    def handle_img(self, current_attrs):
        src = current_attrs['src']
        if src:
            try:
                if isinstance(self.doc, docx.document.Document):
                    self.doc.add_picture(src)
                else:
                    self.add_image_to_cell(self.doc, src)
            except FileNotFoundError:
                src = None

    def handle_table(self):
        """
        To handle nested tables, we will parse tables manually as follows:
        Get table soup
        Create docx table
        Iterate over soup and fill docx table with new instances of this parser
        Tell HTMLParser to ignore any tags until the corresponding closing table tag
        """
        table_soup = self.tables[self.table_no]
        rows, cols = self.get_table_dimensions(table_soup)
        self.table = self.doc.add_table(rows, cols)

        if self.table_style:
            try:
                self.table.style = self.table_style
            except KeyError as e:
                raise ValueError(f"Unable to apply style {self.table_style}.") from e

        rows = self.get_table_rows(table_soup)
        cell_row = 0
        for cell_row, row in enumerate(rows):
            cols = self.get_table_columns(row)
            for cell_col, col in enumerate(cols):
                cell_html = f"<b>{self.get_cell_html(col)}</b>" if col.name == 'th' else self.get_cell_html(col)
                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)
                child_parser.add_html_to_cell(cell_html, self.table.cell(cell_row, cell_col))
        
        self.instances_to_skip = len(table_soup.find_all('table'))
        self.skip_tag = 'table'
        self.skip = True
        self.table = None

    def handle_link(self, href, text):
        rel_id = self.paragraph.part.relate_to(
            href,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True
        )

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), rel_id)

        # Create sub-run
        subrun = self.paragraph.add_run()
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # add default color
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), "0000EE")
        rPr.append(c)

        # add underline
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

        subrun._r.append(rPr)
        subrun._r.text = text

        # Add subrun to hyperlink
        hyperlink.append(subrun._r)

        # Add hyperlink to run
        self.paragraph._p.append(hyperlink)

    def handle_starttag(self, tag, attrs):
        if self.skip:
            return
        if tag == 'head':
            self.skip = True
            self.skip_tag = tag
            self.instances_to_skip = 0
            return
        elif tag == 'body':
            return
        if tag == 'style':
            self.style = True
            return

        current_attrs = dict(attrs)
        if tag == 'span':
            self.tags['span'].append(current_attrs)
            return
        elif tag == 'ol' or tag == 'ul':
            self.tags['list'].append(tag)
            return
        elif tag == 'br':
            self.run.add_break()
            return
        self.tags[tag] = current_attrs
        if tag in ['p', 'pre']:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()
            if self.blockquote:
                self.paragraph.paragraph_format.left_indent = Inches(BLOCKQUOTE_INDENT)
                self.paragraph.paragraph_format.line_spacing = 1       
        elif tag == 'li':
            self.handle_li()
        elif tag == "hr":
            # Horizontal Rule Spceial Handle
            # https://github.com/python-openxml/python-docx/issues/105#issuecomment-62806373
            self.paragraph = self.doc.add_paragraph()
            pPr = self.paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            pPr.insert_element_before(pBdr,
                'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                'w:pPrChange'
            )
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
        elif re.match('h[1-9]', tag):
            if isinstance(self.doc, docx.document.Document):
                h_size = int(tag[1])
                self.paragraph = self.doc.add_heading(level=min(h_size, 9))
            else:
                self.paragraph = self.doc.add_paragraph()
        elif tag == 'img':
            self.handle_img(current_attrs)
            return
        elif tag == 'table':
            self.handle_table()
            return
        elif tag == 'blockquote':
            self.blockquote = True
            return
        # set new run reference point in case of leading line breaks
        if tag in ['p', 'li', 'pre']:
            self.run = self.paragraph.add_run()
        # add style
        if 'style' in current_attrs and self.paragraph:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_paragraph(style)

    def handle_endtag(self, tag):
        if self.skip:
            if not tag == self.skip_tag:
                return
            if self.instances_to_skip > 0:
                self.instances_to_skip -= 1
                return
            self.skip = False
            self.skip_tag = None
            self.paragraph = None

        if tag == 'style':
            self.style = False
            return

        if tag == 'span':
            if self.tags['span']:
                self.tags['span'].pop()
                return
        elif tag == 'ol' or tag == 'ul':
            remove_last_occurence(self.tags['list'], tag)
            # Add a line break after the list
            self.paragraph.add_run().add_break()
            return
        elif tag == 'table':
            self.table_no += 1
            self.table = None
            self.doc = self.document
            self.paragraph = None
        elif tag == 'blockquote':
            self.blockquote = False

        if tag in self.tags:
            self.tags.pop(tag)

    def handle_data(self, data):
        if self.skip or self.style:
            return

        #if 'pre' not in self.tags:
        data = remove_whitespace(data, True, True) if 'pre' not in self.tags else data

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()
            if self.blockquote:
                self.paragraph.paragraph_format.left_indent = Inches(BLOCKQUOTE_INDENT)
                self.paragraph.paragraph_format.line_spacing = 1       

        # There can only be one nested link in a valid html document
        # You cannot have interactive content in an A tag, this includes links
        # https://html.spec.whatwg.org/#interactive-content
        link = self.tags.get('a')
        if link:
            self.handle_link(link['href'], data)
        else:
            self.run = self.paragraph.add_run(data)

            # add font style and name
            for tag in self.tags:
                if tag in font_styles:
                    font_style = font_styles[tag]
                    setattr(self.run.font, font_style, True)

                if tag in font_names:
                    font_name = font_names[tag]
                    self.run.font.name = font_name
                
                if tag == 'code':
                    self.run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            
            spans = self.tags['span']
            for span in spans:
                if 'style' in span:
                    style = self.parse_dict_string(span['style'])
                    self.add_styles_to_run(style)

    def ignore_nested_tables(self, tables_soup):
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables

    def get_table_rows(self, table_soup):
        return table_soup.select(', '.join(self.table_row_selectors), recursive=False)

    def get_table_columns(self, row):
        return row.find_all(['th', 'td'], recursive=False) if row else []

    def get_table_dimensions(self, table_soup):
        rows = self.get_table_rows(table_soup)
        cols = self.get_table_columns(rows[0]) if rows else []
        return len(rows), len(cols)

    def get_tables(self):
        self.tables = self.ignore_nested_tables(self.soup.find_all('table'))  
        self.table_no = 0

    def run_process(self, html):
        self.soup = BeautifulSoup(html, 'html.parser')
        html = str(self.soup)
        self.get_tables()
        self.feed(html)

    def add_html_to_document(self, html, document):
        if not isinstance(html, str):
            raise ValueError('First argument needs to be a %s' % str)
        elif not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.document.Document)
        self.set_initial_attrs(document)
        self.run_process(html)

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.table._Cell)
        unwanted_paragraph = cell.paragraphs[0]
        if unwanted_paragraph.text == "":
            delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')  

    def parse_html_file(self, filename_html, filename_docx=None):
        with open(filename_html, 'r', encoding='utf-8') as infile:
            html = infile.read()
        self.set_initial_attrs()
        self.run_process(html)
        if not filename_docx:
            path, filename = os.path.split(filename_html)
            filename_docx = '%s/new_docx_file_%s' % (path, filename)
        # cleanup empty paragraph at the beginning
        for paragraph in self.doc.paragraphs:
            if len(paragraph.text) == 0:
                delete_paragraph(paragraph)
            else:
                break
        self.doc.save('%s.docx' % filename_docx)
    
    def parse_html_string(self, html):
        self.set_initial_attrs()
        self.run_process(html)
        return self.doc
